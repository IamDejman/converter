"""PDF ↔ DOCX conversion logic.

pdf_to_docx_bytes: Extract text from PDF and create a DOCX.
docx_to_pdf_bytes: Read DOCX paragraphs and render to PDF with Liberation Serif.
"""
import io
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

_FONT_DIR = Path(__file__).resolve().parent.parent / "fonts"


def pdf_to_docx_bytes(pdf_bytes: bytes) -> bytes:
    """Extract text from PDF and create a DOCX document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            if i > 0:
                doc.add_page_break()
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Read DOCX and render to PDF using Liberation Serif."""
    doc = Document(io.BytesIO(docx_bytes))

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_font("Liberation", "", str(_FONT_DIR / "LiberationSerif-Regular.ttf"), uni=True)
    pdf.add_font("Liberation", "B", str(_FONT_DIR / "LiberationSerif-Bold.ttf"), uni=True)
    pdf.add_font("Liberation", "I", str(_FONT_DIR / "LiberationSerif-Italic.ttf"), uni=True)
    pdf.add_font("Liberation", "BI", str(_FONT_DIR / "LiberationSerif-BoldItalic.ttf"), uni=True)
    pdf.add_page()
    pdf.set_font("Liberation", size=12)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            pdf.ln(4)
            continue

        # Detect heading styles
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            pdf.set_font("Liberation", "B", 22)
            pdf.multi_cell(0, 10, text)
            pdf.ln(2)
            pdf.set_font("Liberation", size=12)
        elif "heading 2" in style_name:
            pdf.set_font("Liberation", "B", 18)
            pdf.multi_cell(0, 9, text)
            pdf.ln(2)
            pdf.set_font("Liberation", size=12)
        elif "heading 3" in style_name:
            pdf.set_font("Liberation", "B", 15)
            pdf.multi_cell(0, 8, text)
            pdf.ln(2)
            pdf.set_font("Liberation", size=12)
        else:
            # Check for bold/italic runs
            has_bold = any(r.bold for r in para.runs if r.bold is not None)
            has_italic = any(r.italic for r in para.runs if r.italic is not None)
            if has_bold and has_italic:
                pdf.set_font("Liberation", "BI", 12)
            elif has_bold:
                pdf.set_font("Liberation", "B", 12)
            elif has_italic:
                pdf.set_font("Liberation", "I", 12)
            else:
                pdf.set_font("Liberation", "", 12)
            pdf.multi_cell(0, 7, text)
            pdf.ln(1)
            pdf.set_font("Liberation", "", 12)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()
