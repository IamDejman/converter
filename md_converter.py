"""
Markdown → DOCX / PDF conversion logic.
"""
import io
from html.parser import HTMLParser
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Markdown → HTML
# ---------------------------------------------------------------------------

_MD_EXTENSIONS = [
    "tables",
    "fenced_code",
    "codehilite",
    "toc",
    "nl2br",
    "sane_lists",
    "smarty",
]

_MD_EXT_CONFIGS = {
    "codehilite": {"css_class": "code", "guess_lang": False},
}


def md_to_html(text: str) -> str:
    """Convert markdown text to an HTML fragment."""
    return markdown.markdown(
        text,
        extensions=_MD_EXTENSIONS,
        extension_configs=_MD_EXT_CONFIGS,
    )


# ---------------------------------------------------------------------------
# Styled HTML (for preview & PDF)
# ---------------------------------------------------------------------------

_CSS = """\
body {
  font-family: "Times New Roman", Times, serif;
  font-size: 14px;
  line-height: 1.7;
  color: #1f2328;
  max-width: 800px;
  margin: 0 auto;
  padding: 32px 24px;
}
h1, h2, h3, h4, h5, h6 { margin-top: 1.4em; margin-bottom: 0.6em; font-weight: 600; }
h1 { font-size: 2em; border-bottom: 1px solid #d1d9e0; padding-bottom: 0.3em; }
h2 { font-size: 1.5em; border-bottom: 1px solid #d1d9e0; padding-bottom: 0.3em; }
h3 { font-size: 1.25em; }
p { margin: 0 0 16px; }
a { color: #0969da; text-decoration: none; }
a:hover { text-decoration: underline; }
code {
  background: #eff1f3; padding: 2px 6px; border-radius: 4px;
  font-family: "Times New Roman", Times, serif; font-size: 0.88em;
}
pre {
  background: #f6f8fa; border: 1px solid #d1d9e0; border-radius: 6px;
  padding: 16px; overflow-x: auto; margin: 0 0 16px;
}
pre code { background: none; padding: 0; font-size: 0.88em; }
blockquote {
  border-left: 4px solid #d1d9e0; padding: 0 16px; color: #656d76; margin: 0 0 16px;
}
table { border-collapse: collapse; width: 100%; margin: 0 0 16px; }
th, td { border: 1px solid #d1d9e0; padding: 8px 14px; text-align: left; }
th { background: #f6f8fa; font-weight: 600; }
ul, ol { margin: 0 0 16px; padding-left: 2em; }
li { margin-bottom: 4px; }
hr { border: none; border-top: 2px solid #d1d9e0; margin: 24px 0; }
img { max-width: 100%; }
"""


def md_to_styled_html(text: str) -> str:
    """Return a full HTML document with CSS styling (for PDF / preview)."""
    body = md_to_html(text)
    return (
        "<!DOCTYPE html>\n<html>\n<head>\n"
        '<meta charset="utf-8"/>\n'
        f"<style>{_CSS}</style>\n"
        f"</head>\n<body>\n{body}\n</body>\n</html>"
    )


# ---------------------------------------------------------------------------
# HTML → DOCX builder  (parses the subset of HTML that `markdown` emits)
# ---------------------------------------------------------------------------

class _DocxBuilder(HTMLParser):
    """Walk markdown-generated HTML and build a python-docx Document."""

    def __init__(self):
        super().__init__()
        self.doc = Document()
        # Set default font to Times New Roman
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        for i in range(1, 7):
            hs = self.doc.styles[f"Heading {i}"]
            hs.font.name = "Times New Roman"
        for style_name in ("List Bullet", "List Number"):
            if style_name in self.doc.styles:
                self.doc.styles[style_name].font.name = "Times New Roman"
        self._p = None            # current Paragraph object
        self._bold = False
        self._italic = False
        self._mono = False
        self._in_pre = False
        self._pre_text = ""
        self._list_stack: list[str] = []   # 'ul' | 'ol'
        self._list_counters: list[int] = []
        self._href: str | None = None

        # Table state
        self._in_table = False
        self._table_data: list[list[str]] = []
        self._current_row: list[str] = []
        self._cell_text = ""
        self._is_header_row = False

        # Blockquote
        self._in_blockquote = False

    # -- helpers ----------------------------------------------------------

    def _ensure_paragraph(self):
        if self._p is None:
            self._p = self.doc.add_paragraph()
        return self._p

    def _add_run(self, text: str):
        p = self._ensure_paragraph()
        run = p.add_run(text)
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True
        if self._mono:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)
        return run

    # -- tag handlers -----------------------------------------------------

    def handle_starttag(self, tag, attrs):
        attrs_d = dict(attrs)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            self._p = self.doc.add_heading(level=level)

        elif tag == "p":
            if self._in_table:
                return
            if self._in_blockquote:
                self._p = self.doc.add_paragraph()
                pf = self._p.paragraph_format
                pf.left_indent = Cm(1.5)
            else:
                self._p = self.doc.add_paragraph()

        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True

        elif tag == "code":
            if self._in_pre:
                pass  # handled by <pre>
            else:
                self._mono = True

        elif tag == "pre":
            self._in_pre = True
            self._pre_text = ""

        elif tag in ("ul", "ol"):
            self._list_stack.append(tag)
            self._list_counters.append(0)

        elif tag == "li":
            if self._list_stack:
                kind = self._list_stack[-1]
                if kind == "ul":
                    self._p = self.doc.add_paragraph(style="List Bullet")
                else:
                    self._list_counters[-1] += 1
                    self._p = self.doc.add_paragraph(style="List Number")
            else:
                self._p = self.doc.add_paragraph()

        elif tag == "a":
            self._href = attrs_d.get("href", "")

        elif tag == "blockquote":
            self._in_blockquote = True

        elif tag == "hr":
            p = self.doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xD1, 0xD9, 0xE0)

        elif tag == "br":
            if self._p is not None:
                self._p.add_run("\n")

        # ── Table ───────────────────────────────────────────────
        elif tag == "table":
            self._in_table = True
            self._table_data = []

        elif tag == "thead":
            self._is_header_row = True

        elif tag == "tbody":
            self._is_header_row = False

        elif tag == "tr":
            self._current_row = []

        elif tag in ("td", "th"):
            self._cell_text = ""

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._p = None

        elif tag == "p":
            if not self._in_table:
                self._p = None

        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False

        elif tag == "code":
            if not self._in_pre:
                self._mono = False

        elif tag == "pre":
            self._in_pre = False
            # Add code block as a styled paragraph
            p = self.doc.add_paragraph()
            run = p.add_run(self._pre_text.rstrip("\n"))
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(4)
            # Light grey shading via XML
            shading = self.doc.element.makeelement(
                qn("w:shd"),
                {qn("w:fill"): "F6F8FA", qn("w:val"): "clear"},
            )
            p._element.get_or_add_pPr().append(shading)
            self._pre_text = ""

        elif tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
                self._list_counters.pop()

        elif tag == "li":
            self._p = None

        elif tag == "a":
            self._href = None

        elif tag == "blockquote":
            self._in_blockquote = False

        # ── Table ───────────────────────────────────────────────
        elif tag in ("td", "th"):
            self._current_row.append(self._cell_text.strip())

        elif tag == "tr":
            self._table_data.append(self._current_row)

        elif tag == "thead":
            self._is_header_row = False

        elif tag == "table":
            self._in_table = False
            self._build_table()

    def handle_data(self, data):
        # Accumulate pre text
        if self._in_pre:
            self._pre_text += data
            return

        # Table cell text
        if self._in_table:
            self._cell_text += data
            return

        # Normal inline text
        if data.strip() == "" and self._p is None:
            return

        self._add_run(data)

    def _build_table(self):
        """Build a docx table from collected _table_data."""
        if not self._table_data:
            return
        num_cols = max(len(row) for row in self._table_data)
        num_rows = len(self._table_data)
        tbl = self.doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

        for r, row_data in enumerate(self._table_data):
            for c, cell_text in enumerate(row_data):
                cell = tbl.cell(r, c)
                cell.text = cell_text
                # Bold the header row
                if r == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        self._table_data = []


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def md_to_docx_bytes(text: str) -> bytes:
    """Convert markdown text to DOCX, return raw bytes."""
    builder = _DocxBuilder()
    html = md_to_html(text)
    builder.feed(html)
    buf = io.BytesIO()
    builder.doc.save(buf)
    buf.seek(0)
    return buf.read()


_FONT_DIR = Path(__file__).parent / "fonts"


def md_to_pdf_bytes(text: str) -> bytes:
    """Convert markdown text to PDF, return raw bytes."""
    from fpdf import FPDF
    from html.parser import HTMLParser as _HP

    html = md_to_html(text)

    class _PdfBuilder(_HP):
        """Minimal HTML->fpdf2 renderer for markdown-generated HTML."""

        def __init__(self):
            super().__init__()
            self.pdf = FPDF()
            self.pdf.set_auto_page_break(auto=True, margin=20)
            # Register Liberation Serif (Unicode-capable, metrically identical to Times New Roman)
            self.pdf.add_font("Liberation", "", str(_FONT_DIR / "LiberationSerif-Regular.ttf"), uni=True)
            self.pdf.add_font("Liberation", "B", str(_FONT_DIR / "LiberationSerif-Bold.ttf"), uni=True)
            self.pdf.add_font("Liberation", "I", str(_FONT_DIR / "LiberationSerif-Italic.ttf"), uni=True)
            self.pdf.add_font("Liberation", "BI", str(_FONT_DIR / "LiberationSerif-BoldItalic.ttf"), uni=True)
            self.pdf.add_page()
            self.pdf.set_font("Liberation", size=12)
            self._bold = False
            self._italic = False
            self._in_pre = False
            self._pre_text = ""
            self._in_li = False
            self._list_stack: list[str] = []
            self._list_counters: list[int] = []
            self._heading_level = 0
            self._in_table = False
            self._table_data: list[list[str]] = []
            self._current_row: list[str] = []
            self._cell_text = ""

        def _set_font(self):
            style = ""
            if self._bold:
                style += "B"
            if self._italic:
                style += "I"
            size = 12
            if self._heading_level:
                size = {1: 24, 2: 20, 3: 16, 4: 14, 5: 12, 6: 11}.get(self._heading_level, 12)
            if self._in_pre:
                self.pdf.set_font("Courier", style=style, size=9)
            else:
                self.pdf.set_font("Liberation", style=style, size=size)

        def handle_starttag(self, tag, attrs):
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self._heading_level = int(tag[1])
                self._bold = True
                self._set_font()
                self.pdf.ln(4)
            elif tag in ("strong", "b"):
                self._bold = True
                self._set_font()
            elif tag in ("em", "i"):
                self._italic = True
                self._set_font()
            elif tag == "pre":
                self._in_pre = True
                self._pre_text = ""
            elif tag == "code" and not self._in_pre:
                self.pdf.set_font("Courier", size=10)
            elif tag in ("ul", "ol"):
                self._list_stack.append(tag)
                self._list_counters.append(0)
            elif tag == "li":
                self._in_li = True
                if self._list_stack:
                    kind = self._list_stack[-1]
                    if kind == "ul":
                        self.pdf.cell(10)
                        self.pdf.cell(6, 6, "\u2022 ")
                    else:
                        self._list_counters[-1] += 1
                        self.pdf.cell(10)
                        self.pdf.cell(6, 6, f"{self._list_counters[-1]}. ")
            elif tag == "br":
                self.pdf.ln(5)
            elif tag == "hr":
                self.pdf.ln(4)
                self.pdf.set_draw_color(209, 217, 224)
                self.pdf.line(10, self.pdf.get_y(), 200, self.pdf.get_y())
                self.pdf.ln(4)
            elif tag == "table":
                self._in_table = True
                self._table_data = []
            elif tag == "tr":
                self._current_row = []
            elif tag in ("td", "th"):
                self._cell_text = ""

        def handle_endtag(self, tag):
            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                self.pdf.ln(6)
                self._heading_level = 0
                self._bold = False
                self._set_font()
            elif tag == "p":
                if not self._in_table:
                    self.pdf.ln(6)
            elif tag in ("strong", "b"):
                self._bold = False
                self._set_font()
            elif tag in ("em", "i"):
                self._italic = False
                self._set_font()
            elif tag == "pre":
                self._in_pre = False
                self._set_font()
                self.pdf.set_fill_color(246, 248, 250)
                self.pdf.set_font("Courier", size=9)
                for line in self._pre_text.rstrip("\n").split("\n"):
                    self.pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT", fill=True)
                self.pdf.ln(4)
                self._set_font()
                self._pre_text = ""
            elif tag == "code" and not self._in_pre:
                self._set_font()
            elif tag in ("ul", "ol"):
                if self._list_stack:
                    self._list_stack.pop()
                    self._list_counters.pop()
                self.pdf.ln(3)
            elif tag == "li":
                self._in_li = False
                self.pdf.ln(5)
            elif tag in ("td", "th"):
                self._current_row.append(self._cell_text.strip())
            elif tag == "tr":
                self._table_data.append(self._current_row)
            elif tag == "table":
                self._in_table = False
                self._build_table()

        def handle_data(self, data):
            if self._in_pre:
                self._pre_text += data
                return
            if self._in_table:
                self._cell_text += data
                return
            if data.strip() == "":
                return
            self.pdf.write(6, data)

        def _build_table(self):
            if not self._table_data:
                return
            num_cols = max(len(r) for r in self._table_data)
            if num_cols == 0:
                return
            col_w = (self.pdf.w - 20) / num_cols
            self.pdf.set_font("Liberation", size=10)
            for r, row in enumerate(self._table_data):
                for c, cell in enumerate(row):
                    if r == 0:
                        self.pdf.set_font("Liberation", style="B", size=10)
                    else:
                        self.pdf.set_font("Liberation", size=10)
                    self.pdf.cell(col_w, 7, cell, border=1)
                self.pdf.ln()
            self.pdf.ln(4)
            self._set_font()
            self._table_data = []

    builder = _PdfBuilder()
    builder.feed(html)
    return bytes(builder.pdf.output())
